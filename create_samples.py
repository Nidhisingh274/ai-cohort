from docx import Document
from reportlab.pdfgen import canvas
from PIL import Image, ImageDraw

# =====================
# 1. Create Word Document
# =====================
doc = Document()
doc.add_heading('Claims Process Guide', 0)
doc.add_paragraph('This document explains how to file a claim.')

doc.add_heading('Step 1: Submit Your Claim', level=1)
doc.add_paragraph('Claims must be submitted within 90 days of the service date. Include your member ID and procedure details.')

doc.add_heading('Step 2: Claim Review', level=1)
doc.add_paragraph('Our team reviews claims within 10 business days. You will receive a status update by email.')

doc.add_heading('Step 3: Reimbursement', level=1)
doc.add_paragraph('Approved claims are reimbursed within 15 business days to your registered payment method.')

doc.save('claims_process_sample.docx')
print("Created: claims_process_sample.docx")

# =====================
# 2. Create PDF
# =====================
c = canvas.Canvas("benefits_sample.pdf")
c.drawString(100, 750, "Summary of Benefits and Coverage")
c.drawString(100, 720, "Plan: Gold PPO")
c.drawString(100, 700, "Monthly Premium: $500")
c.drawString(100, 680, "Annual Deductible: $2000")
c.drawString(100, 660, "Copay: 10% after deductible")
c.drawString(100, 640, "Coverage includes preventive care, hospital visits, and prescriptions.")
c.drawString(100, 620, "Network: Nationwide PPO Network")
c.drawString(100, 600, "Out-of-pocket maximum: $5000 per year")
c.save()
print("Created: benefits_sample.pdf")

# =====================
# 3. Create Scanned form image
# =====================
img = Image.new('RGB', (600, 400), color='white')
d = ImageDraw.Draw(img)
d.text((20, 20), "ENROLLMENT FORM", fill='black')
d.text((20, 60), "Member Name: John Doe", fill='black')
d.text((20, 100), "Member ID: M1001", fill='black')
d.text((20, 140), "Plan Selected: Gold PPO", fill='black')
d.text((20, 180), "Date of Birth: 1990-05-15", fill='black')
d.text((20, 220), "Enrollment Date: 2023-01-15", fill='black')
d.text((20, 260), "Coverage Type: PPO", fill='black')
d.text((20, 300), "Signature: John Doe", fill='black')
img.save('enrollment_sample.png')
print("Created: enrollment_sample.png")

print("\nAll sample documents created successfully!")